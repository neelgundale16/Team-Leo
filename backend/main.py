# main.py

import json
import time
import uuid
import logging
import os
from io import BytesIO
from contextlib import asynccontextmanager
from typing import AsyncGenerator, Optional

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse

from models import ChatRequest, StreamToken, SessionStats
from vault import vault, load_demo_financial_data
from sentinel import sentinel
from rewriter import rewriter
from interceptor import stream_and_detect

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s"
)
logger = logging.getLogger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("🛡️  Project Veracity — starting up...")

    vault.initialize()
    if vault.get_count() == 0:
        logger.info("Vault empty — loading demo financial data...")
        load_demo_financial_data(vault)
    logger.info(f"Vault ready | {vault.get_count()} documents loaded")

    sentinel.initialize()
    logger.info("Sentinel ready")

    logger.info("✅ All systems go. Hallucination firewall is active.")
    yield
    logger.info("Project Veracity shutting down.")


app = FastAPI(title="Project Veracity", version="1.0.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)


def sse(event_type: str, data: dict) -> str:
    return f"data: {json.dumps({'event_type': event_type, 'data': data})}\n\n"


async def firewall_generator(query: str) -> AsyncGenerator[str, None]:
    stats = SessionStats()
    pipeline_start = time.perf_counter()

    try:
        async for raw, claims in stream_and_detect(query):

            if raw.startswith("TOKEN:"):
                token_text = raw[6:]
                token_event = StreamToken(
                    id=str(uuid.uuid4())[:8],
                    text=token_text,
                    status="streaming"
                )
                yield sse("token", token_event.dict())
                continue

            if raw.startswith("ERROR:"):
                yield sse("error", {"message": raw[6:].strip()})
                return

            if not claims:
                stats.claims_skipped += 1
                continue

            stats.total_claims_detected += len(claims)

            for claim in claims:
                t0 = time.perf_counter()

                if not sentinel.is_fact_seeking(claim.sentence):
                    stats.claims_skipped += 1
                    continue

                vault_result = vault.search(claim.text)
                if vault_result is None:
                    stats.claims_verified += 1
                    yield sse("stats", stats.dict())
                    continue

                nli_result = sentinel.classify(
                    claim_text=claim.text,
                    context_text=vault_result.matched_text
                )

                latency = (time.perf_counter() - t0) * 1000
                logger.info(
                    f"'{claim.text}' | {nli_result.label} "
                    f"({nli_result.confidence:.2f}) | {latency:.1f}ms"
                )

                n = stats.claims_verified + 1
                stats.avg_verification_latency_ms = (
                    (stats.avg_verification_latency_ms * (n - 1) + latency) / n
                )
                stats.claims_verified += 1

                if nli_result.is_hallucination:
                    stats.hallucinations_found += 1
                    corrected = rewriter.rewrite(
                        claim.sentence,
                        vault_result,
                        nli_result
                    )
                    payload = rewriter.build_correction_payload(
                        claim.sentence,
                        corrected,
                        vault_result.source_document
                    )
                    stats.corrections_made += 1

                    logger.info(
                        f"CORRECTED | '{claim.text}' → {vault_result.source_document}"
                    )

                    yield sse("correction", {
                        "original_claim": claim.text,
                        "original_sentence": claim.sentence,
                        "corrected_sentence": corrected,
                        "source": vault_result.source_document,
                        "similarity_score": vault_result.similarity_score,
                        "nli_label": nli_result.label,
                        "nli_confidence": nli_result.confidence,
                        "diff_ratio": payload.get("diff_ratio", 0),
                    })

                yield sse("stats", stats.dict())

        stats.total_pipeline_latency_ms = (time.perf_counter() - pipeline_start) * 1000
        logger.info(
            f"Done | claims:{stats.total_claims_detected} "
            f"corrections:{stats.corrections_made} "
            f"total:{stats.total_pipeline_latency_ms:.0f}ms"
        )
        yield sse("done", stats.dict())

    except Exception as e:
        logger.error(f"Pipeline error: {e}", exc_info=True)
        yield sse("error", {"message": str(e)})


def chunk_text(text: str, chunk_size: int = 500) -> list[str]:
    words = text.split()
    chunks = []
    chunk = []
    chars = 0

    for word in words:
        chunk.append(word)
        chars += len(word) + 1
        if chars >= chunk_size:
            chunks.append(" ".join(chunk))
            chunk = []
            chars = 0

    if chunk:
        chunks.append(" ".join(chunk))

    return chunks


def extract_text_from_pdf(raw_bytes: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        raise HTTPException(status_code=500, detail="Run: pip install pypdf")

    try:
        reader = pypdf.PdfReader(BytesIO(raw_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        return "\n".join(pages).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF parse error: {e}")


def extract_text_from_docx(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="Run: pip install python-docx")

    try:
        doc = Document(BytesIO(raw_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX parse error: {e}")


def extract_text_from_txt(raw_bytes: bytes) -> str:
    try:
        return raw_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="Text file must be UTF-8 encoded")


@app.post("/chat")
async def chat(request: ChatRequest):
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    logger.info(f"Query: '{request.query[:80]}'")
    return StreamingResponse(
        firewall_generator(request.query),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Access-Control-Allow-Origin": "*",
        }
    )


@app.post("/vault/upload")
async def vault_upload(
    file: UploadFile = File(...),
    source_name: Optional[str] = Form(None)
):
    filename = source_name or file.filename or "uploaded_document"
    content_type = file.content_type or ""
    raw_bytes = await file.read()
    chunks: list[str] = []

    if "pdf" in content_type or filename.lower().endswith(".pdf"):
        try:
            import io
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    words, chunk, chars = text.split(), [], 0
                    for w in words:
                        chunk.append(w)
                        chars += len(w) + 1
                        if chars >= 500:
                            chunks.append(" ".join(chunk))
                            chunk, chars = [], 0
                    if chunk:
                        chunks.append(" ".join(chunk))
        except ImportError:
            raise HTTPException(status_code=500, detail="Run: pip install pypdf")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF parse error: {e}")

    else:
        try:
            text = raw_bytes.decode("utf-8")
            words, chunk, chars = text.split(), [], 0
            for w in words:
                chunk.append(w)
                chars += len(w) + 1
                if chars >= 500:
                    chunks.append(" ".join(chunk))
                    chunk, chars = [], 0
            if chunk:
                chunks.append(" ".join(chunk))
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail="File must be PDF or UTF-8 text")

    if not chunks:
        raise HTTPException(status_code=400, detail="No text extracted from file")

    chunks = chunk_text(text, chunk_size=500)
    if not chunks:
        raise HTTPException(status_code=400, detail="No usable chunks created from file")

    vault.add_documents_bulk(chunks, filename)
    logger.info(f"Uploaded '{filename}' | {len(chunks)} chunks → vault")

    return JSONResponse({
        "status": "success",
        "filename": filename,
        "chunks_added": len(chunks),
        "vault_total": vault.get_count()
    })


@app.get("/health")
async def health():
    return {
        "status": "ok",
        "llm_provider": "Ollama",
        "llm_model": "llama3",
        "vault_documents": vault.get_count(),
        "sentinel_loaded": sentinel._initialized,
        "mock_mode": os.getenv("USE_MOCK", "false"),
        "version": "1.0.0"
    }


@app.get("/vault/count")
async def vault_count():
    return {"count": vault.get_count()}


@app.post("/vault/add")
async def vault_add(payload: dict):
    text = payload.get("text", "").strip()
    source = payload.get("source", "manual_entry")

    if not text:
        raise HTTPException(status_code=400, detail="text field required")

    vault.add_document(text=text, source_name=source)
    return {"status": "added", "vault_total": vault.get_count()}


@app.delete("/vault/clear")
async def vault_clear():
    vault.clear_vault()
    load_demo_financial_data(vault)
    return {"status": "reset", "vault_total": vault.get_count()}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)